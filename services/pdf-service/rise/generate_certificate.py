from docx import Document
import fitz  # PyMuPDF
from typing import Dict

# ISO Standards Mapping - Convert short names to full versions with years
ISO_STANDARDS_MAPPING = {
    # Quality & Management
    "ISO 9001": "ISO 9001:2015",
    "9001": "ISO 9001:2015",
    "ISO 14001": "ISO 14001:2015", 
    "14001": "ISO 14001:2015",
    "ISO 45001": "ISO 45001:2018",
    "45001": "ISO 45001:2018",
    "ISO 50001": "ISO 50001:2018",
    "50001": "ISO 50001:2018",
    "ISO 31000": "ISO 31000:2018",
    "31000": "ISO 31000:2018",
    
    # Food Safety
    "ISO 22000": "ISO 22000:2018",
    "22000": "ISO 22000:2018",
    "ISO/TS 22002-1": "ISO/TS 22002-1:2009",
    "22002-1": "ISO/TS 22002-1:2009",
    "ISO 22005": "ISO 22005:2007",
    "22005": "ISO 22005:2007",
    
    # Laboratory & Testing
    "ISO/IEC 17025": "ISO/IEC 17025:2017",
    "17025": "ISO/IEC 17025:2017",
    "ISO 15189": "ISO 15189:2022",
    "15189": "ISO 15189:2022",
    
    # Information Security & IT
    "ISO/IEC 27001": "ISO/IEC 27001:2022",
    "27001": "ISO/IEC 27001:2022",
    "ISO/IEC 27002": "ISO/IEC 27002:2022",
    "27002": "ISO/IEC 27002:2022",
    "ISO/IEC 20000-1": "ISO/IEC 20000-1:2018",
    "20000-1": "ISO/IEC 20000-1:2018",
    "ISO/IEC 22301": "ISO/IEC 22301:2019",
    "22301": "ISO/IEC 22301:2019",
    
    # Manufacturing & Industrial
    "ISO 13485": "ISO 13485:2016",
    "13485": "ISO 13485:2016",
    "IATF 16949": "IATF 16949:2016",
    "16949": "IATF 16949:2016",
    "ISO 3834-2": "ISO 3834-2:2021",
    "3834-2": "ISO 3834-2:2021",
    
    # Environment & Sustainability
    "ISO 14064-1": "ISO 14064-1:2018",
    "14064-1": "ISO 14064-1:2018",
    "ISO 14046": "ISO 14046:2014",
    "14046": "ISO 14046:2014",
    "ISO 20121": "ISO 20121:2012",
    "20121": "ISO 20121:2012",
    
    # Asset, Facility, and Supply Chain
    "ISO 55001": "ISO 55001:2014",
    "55001": "ISO 55001:2014",
    "ISO 28000": "ISO 28000:2022",
    "28000": "ISO 28000:2022",
    
    # Aerospace
    "AS 9100D": "AS 9100D:2016",
    "9100D": "AS 9100D:2016",
    
    # Other Notable Standards
    "ISO 37001": "ISO 37001:2016",
    "37001": "ISO 37001:2016",
    "ISO 19600": "ISO 19600:2014",
    "19600": "ISO 19600:2014",
    "ISO 29993": "ISO 29993:2017",
    "29993": "ISO 29993:2017",
}

# ISO Standards Code Mapping - For certification codes
ISO_STANDARDS_CODES = {
    "ISO 9001:2015": "CM-MS-7842",
    "ISO 14001:2015": "CM-MS-7836", 
    "ISO 45001:2018": "CM-MS-7832",
    "ISO 22000:2018": "CM-MS-7822",
    "ISO/IEC 27001:2022": "CM-MS-7820",
    "ISO 37001:2016": "CM-MS-7804",
    "ISO/IEC 22301:2019": "CM-MS-7807",
    "ISO 50001:2018": "CM-MS-7814",
    "ISO 20001:2018": "CM-MS-7811",
}

# ISO Standards Descriptions Mapping - For separate use
ISO_STANDARDS_DESCRIPTIONS = {
    "ISO 9001:2015": "Quality Management System",
    "ISO 14001:2015": "Environmental Management System",
    "ISO 45001:2018": "Occupational Health & Safety",
    "ISO 50001:2018": "Energy Management System",
    "ISO 31000:2018": "Risk Management Guidelines",
    "ISO 22000:2018": "Food Safety Management System",
    "ISO/TS 22002-1:2009": "Prerequisite programs on food safety",
    "ISO 22005:2007": "Traceability in the feed and food chain",
    "ISO/IEC 17025:2017": "Testing and Calibration Laboratories",
    "ISO 15189:2022": "Medical Laboratories – Quality and Competence",
    "ISO/IEC 27001:2022": "Information Security Management System",
    "ISO/IEC 27002:2022": "Information Security Controls",
    "ISO/IEC 20000-1:2018": "IT Service Management System",
    "ISO/IEC 22301:2019": "Business Continuity Management System",
    "ISO 13485:2016": "Medical Devices – Quality Management System",
    "IATF 16949:2016": "Automotive Quality Management System",
    "ISO 3834-2:2021": "Quality requirements for fusion welding",
    "ISO 14064-1:2018": "Greenhouse Gases",
    "ISO 14046:2014": "Water Footprint",
    "ISO 20121:2012": "Event Sustainability Management System",
    "ISO 55001:2014": "Asset Management System",
    "ISO 28000:2022": "Security Management Systems for Supply Chain",
    "AS 9100D:2016": "Aerospace Quality (based on ISO 9001:2015)",
    "ISO 37001:2016": "Anti-bribery Management System",
    "ISO 19600:2014": "Compliance Management System",
    "ISO 29993:2017": "Learning Services",
}

def expand_iso_standard(iso_text: str) -> str:
    """Expand ISO standard name to full version with year if available."""
    if not iso_text:
        return iso_text
    
    # Clean the input text
    cleaned_text = iso_text.strip()
    
    # First, try exact match
    if cleaned_text in ISO_STANDARDS_MAPPING:
        return ISO_STANDARDS_MAPPING[cleaned_text]
    
    # If no exact match, try to find partial matches
    # This handles cases where users might enter variations
    for short_name, full_name in ISO_STANDARDS_MAPPING.items():
        # Check if the input contains the standard number
        if short_name.lower() in cleaned_text.lower():
            return full_name
    
    # If still no match, try to extract just the number and match
    # This handles cases like "37001" when we have "37001" in mapping
    import re
    number_match = re.search(r'(\d+(?:-\d+)?)', cleaned_text)
    if number_match:
        number = number_match.group(1)
        if number in ISO_STANDARDS_MAPPING:
            return ISO_STANDARDS_MAPPING[number]
    
    # If no match found, return original text
    return iso_text

def get_iso_standard_code(iso_standard: str) -> str:
    """
    Get the certification code for a given ISO standard.
    Example: "ISO 9001:2015" -> "CM-MS-7842"
    Returns empty string if no code mapping exists.
    """
    if not iso_standard:
        return ""
    
    # First expand the ISO standard if it's a short name
    expanded_iso = expand_iso_standard(iso_standard)
    
    # Look up the code in the mapping
    return ISO_STANDARDS_CODES.get(expanded_iso, "")

def parse_word_form(docx_path: str) -> Dict[str, str]:
    """Parse the first table in a Word document and extract required fields."""
    
    doc = Document(docx_path)
    
    if len(doc.tables) == 0:
        raise Exception("No tables found in document")
    
    table = doc.tables[0]
    
    data = {}
    for i, row in enumerate(table.rows):
        if len(row.cells) == 2:
            key = row.cells[0].text.strip()
            value = row.cells[1].text.strip()
            data[key] = value
        else:
            continue
    
    # Get ISO Standard and expand it to full version with year
    iso_standard = data.get("ISO Standard Required", "").splitlines()[-1]
    expanded_iso = expand_iso_standard(iso_standard)
    
    result = {
        "Company Name": data.get("Company Name", ""),
        "Address": data.get("Address", ""),
        "ISO Standard": expanded_iso,
        "Scope": data.get("Scope", "")
    }
    
    return result

def get_text_height(text: str, fontsize: float, fontname: str, max_width: float) -> float:
    """Estimate the height of a text block when wrapped to fit max_width."""
    font = fitz.Font(fontname=fontname)
    words = text.split()
    lines = []
    current_line = ""
    for word in words:
        test_line = current_line + (" " if current_line else "") + word
        if font.text_length(test_line, fontsize) <= max_width:
            current_line = test_line
        else:
            lines.append(current_line)
            current_line = word
    if current_line:
        lines.append(current_line)
    return len(lines) * fontsize * 1.2  # Approximate line height with spacing

def insert_centered_textbox(
    page: fitz.Page,
    rect: fitz.Rect,
    text: str,
    fontname: str,
    fontsize: float,
    color: tuple
) -> None:
    """Insert text centered both vertically and horizontally in the given rectangle."""
    # Calculate text height and center it vertically
    text_height = get_text_height(text, fontsize, fontname, rect.width)
    start_y = rect.y0 + (rect.height - text_height) / 2
    box = fitz.Rect(rect.x0, start_y, rect.x1, start_y + text_height)
    
    page.insert_textbox(
        box,
        text,
        fontsize=fontsize,
        fontname=fontname,
        color=color,
        align=1  # Centered
    )

def generate_certificate(base_pdf_path: str, output_pdf_path: str, values: Dict[str, str], template_type: str = "standard") -> Dict[str, any]:
    """Generate a certificate PDF by overlaying extracted values onto a template.
    
    Returns:
        Dict containing success status and overflow warnings
    """

    
    # Initialize tracking for overflow warnings
    overflow_warnings = []
    doc = fitz.open(base_pdf_path)
    page = doc[0]

    # --- Configuration ---
    color = (0, 0, 0)  # Black text
    fontname = "Times-Bold"  # Use Times New Roman Bold font
    
    # ✅ ADDED: Font weight preservation system
    def detect_font_weight(text):
        """
        Detect font weight from text formatting.
        Excel doesn't preserve bold formatting, but we can implement
        a marker system for future enhancement.
        """
        # For now, return default font
        # Future enhancement: Parse Excel formatting or use markers like **bold** or __bold__
        return "Times-Bold"
    
    def process_bold_text(text):
        """
        Process text with bold markers and return segments with font information.
        Returns list of tuples: (text_segment, font_name, is_bold)
        """
        if not text:
            return [(text, "Times-Bold", False)]
        
        segments = []
        current_text = text
        
        # Process **bold** markers
        while '**' in current_text:
            parts = current_text.split('**', 2)
            if len(parts) >= 3:
                # Add normal text before bold
                if parts[0]:
                    segments.append((parts[0], "Times-Roman", False))
                # Add bold text
                segments.append((parts[1], "Times-Bold", True))
                # Continue with remaining text
                current_text = parts[2]
            else:
                break
        
        # Process __bold__ markers
        while '__' in current_text:
            parts = current_text.split('__', 2)
            if len(parts) >= 3:
                # Add normal text before bold
                if parts[0]:
                    segments.append((parts[0], "Times-Roman", False))
                # Add bold text
                segments.append((parts[1], "Times-Bold", True))
                # Continue with remaining text
                current_text = parts[2]
            else:
                break
        
        # Add any remaining normal text
        if current_text:
            segments.append((current_text, "Times-Roman", False))
        
        # If no bold markers found, return original text with default font
        if not segments:
            segments.append((text, "Times-Bold", False))
        
        return segments

    def get_font_for_text(text, default_font="Times-Bold"):
        """
        Get appropriate font for text based on content analysis.
        This is now a legacy function - use process_bold_text for full processing.
        """
        # Check for bold markers
        if '**' in text or '__' in text:
            return "Times-Bold"  # Will be processed by process_bold_text
        else:
            return default_font

    def render_mixed_format_text(page, position, text, font_size, color, max_width=None):
        """
        Render text with mixed bold/normal formatting at the specified position.
        Returns the total width used for positioning calculations.
        """
        if not text:
            return 0
        
        segments = process_bold_text(text)
        current_x = position[0]
        total_width = 0
        
        for segment_text, font_name, is_bold in segments:
            if not segment_text:
                continue
                
            # Calculate text width
            font_obj = fitz.Font(fontname=font_name)
            text_width = font_obj.text_length(segment_text, font_size)
            
            # Check if we need to wrap (if max_width is specified)
            if max_width and current_x + text_width > position[0] + max_width:
                # For now, just render what fits - could implement word wrapping here
                pass
            
            # Render the text segment
            page.insert_text(
                (current_x, position[1]),
                segment_text,
                fontsize=font_size,
                fontname=font_name,
                color=color
            )
            
            # Move position for next segment
            current_x += text_width
            total_width += text_width
        
        return total_width

    # ✅ ADDED: Logo processing
    logo_lookup = values.get("logo_lookup", {})
    logo_filename = values.get("Logo", "").strip()
    
    # Process logo if specified and available
    logo_image = None
    if logo_filename and logo_lookup and logo_filename in logo_lookup:
        try:
            # Convert uploaded file to PIL Image
            from PIL import Image
            import io
            
            logo_file = logo_lookup[logo_filename]
            if hasattr(logo_file, 'file'):
                # Reset file pointer
                logo_file.file.seek(0)
                # Read file content
                logo_content = logo_file.file.read()
                # Convert to PIL Image
                logo_image = Image.open(io.BytesIO(logo_content))
            else:
                logo_image = None
        except Exception as logo_error:
            logo_image = None
    else:
        logo_image = None
    
    # Standard template coordinates (current)
    standard_coords = {
        "management_system": fitz.Rect(87.9, 185, 580, 226.6),
        "Company Name and Address": fitz.Rect(87.9, 222.6, 580, 315),
        "ISO Standard": fitz.Rect(194.9, 334, 460.3, 370),
        "Scope": {
            "short": fitz.Rect(87.9, 386, 580, 475),    # <24 lines
            "long": fitz.Rect(87.9, 373, 580, 486)      # 24-30 lines
        },
        "certification_code": fitz.Rect(253, 757, 285, 762)  # ✅ UPDATED: Certification code coordinates
    }
    
    # Large template coordinates (for >11 lines)
    # All generation types use Y0=354
    large_coords = {
        "management_system": fitz.Rect(87.9, 185, 580, 226.6),  # Same as standard
        "Company Name and Address": fitz.Rect(87.9, 222.6, 580, 295),  # Same as standard
        "ISO Standard": fitz.Rect(194.9, 300, 460.3, 336),  # Same as standard
        "Scope": fitz.Rect(85, 354, 577, 536),  # Y0=354 for all generation types
        "certification_code": fitz.Rect(253, 757, 285, 762)  # ✅ UPDATED: Certification code coordinates
    }
    
    # Logo template coordinates (when Logo = "yes")
    logo_coords = {
        "management_system": fitz.Rect(87.9, 185, 580, 226.6),  # Same as standard
        "logo": fitz.Rect(87.9, 226.6, 580, 262.6),  # Logo area: below management_system, above company name
        "Company Name and Address": fitz.Rect(87.9, 262.6, 580, 355),  # Lowered y-coordinates
        "ISO Standard": fitz.Rect(194.9, 374, 460.3, 410),  # Lowered y-coordinates
        "Scope": {
            "short": fitz.Rect(87.9, 426, 580, 515),    # Lowered y-coordinates, y2 remains same
            "long": fitz.Rect(87.9, 413, 580, 526)      # Lowered y-coordinates, y2 remains same
        },
        "certification_code": fitz.Rect(253, 757, 285, 762)  # ✅ UPDATED: Certification code coordinates
    }
    
    # ✅ ADDED: Defensive checks for coordinate dictionaries
    if standard_coords is None:
        print(f"⚠️ [CERTIFICATE] standard_coords is None - cannot continue")
        raise ValueError("standard_coords is None - cannot generate certificate")
    
    if large_coords is None:
        print(f"⚠️ [CERTIFICATE] large_coords is None - cannot continue")
        raise ValueError("large_coords is None - cannot generate certificate")
    
    if logo_coords is None:
        print(f"⚠️ [CERTIFICATE] logo_coords is None - cannot continue")
        raise ValueError("logo_coords is None - cannot generate certificate")
    
    # Select coordinates based on template type
    if template_type in ["standard", "standard_eco", "standard_nonaccredited"]:
        coords = standard_coords
    elif template_type in ["large", "large_eco", "large_nonaccredited"]:
        coords = large_coords
    elif template_type in ["logo", "logo_nonaccredited", "logo_other", "logo_nonaccredited_other"]:
        coords = logo_coords
    else:
        # Default to standard coordinates
        coords = standard_coords
    
    # ✅ ADDED: Defensive check for selected coords
    if coords is None:
        raise ValueError("Selected coords is None - cannot generate certificate")

    
    font_starts = {
        "Company Name and Address": 45,  # Company Name starts from 45pt
        "Scope": 20,
        "ISO Standard": 80,
        "management_system": 15,  # Management system line font size
    }
    # --- End Configuration ---

    # Determine Scope coordinates based on content length
    scope_text = values.get("Scope", "")
    scope_words = len(scope_text.split())
    
    # Calculate estimated lines for Scope (approximate calculation)
    estimated_lines = max(1, (scope_words * 8) // 60)  # Rough estimate: 8 chars per word, 60 chars per line
    
    # ✅ ADDED: Defensive check for Scope coordinates
    if "Scope" not in coords:
        print(f"⚠️ [CERTIFICATE] Scope coordinates not found in coords - cannot continue")
        raise ValueError("Scope coordinates not found - cannot generate certificate")
    
    # Determine which coordinate set to use (lines win over words)
    if template_type in ["standard", "standard_eco", "standard_nonaccredited", "logo", "logo_nonaccredited", "logo_other", "logo_nonaccredited_other"]:
        # Standard template: dynamic coordinates based on content length
        if estimated_lines >= 24:  # Long content condition
            if "long" not in coords["Scope"]:
                print(f"⚠️ [CERTIFICATE] Scope long coordinates not found - cannot continue")
                raise ValueError("Scope long coordinates not found - cannot generate certificate")
            scope_rect = coords["Scope"]["long"]
            scope_layout = "long"
            print(f"🎯 [CERTIFICATE] Scope: {estimated_lines} lines (≥24) -> selected LONG scope coordinates")

        else:  # Short content condition
            if "short" not in coords["Scope"]:
                print(f"⚠️ [CERTIFICATE] Scope short coordinates not found - cannot continue")
                raise ValueError("Scope short coordinates not found - cannot generate certificate")
            scope_rect = coords["Scope"]["short"]
            scope_layout = "short"
            print(f"🎯 [CERTIFICATE] Scope: {estimated_lines} lines (<24) -> selected SHORT scope coordinates")

    else:
        # Large template: fixed large coordinates
        scope_rect = coords["Scope"]
        scope_layout = "large"
        print(f"🎯 [CERTIFICATE] Scope: {estimated_lines} lines -> selected LARGE scope coordinates (fixed)")
    
    print(f"🎯 [CERTIFICATE] Final scope coordinates: {scope_rect}")

    # Store original scope coordinates before modification (for Extra Line processing)
    original_scope_coords = coords["Scope"].copy() if isinstance(coords["Scope"], dict) else coords["Scope"]
    
    # Add Scope coordinates to the main coords dictionary
    coords["Scope"] = scope_rect
    
    # ✅ NEW: Adjust scope coordinates when Extra Line is present
    extra_line = values.get("Extra Line", "").strip()
    if extra_line:
        print(f"🔍 [CERTIFICATE] Extra Line present - adjusting scope coordinates")
        
        # Adjust scope coordinates: reduce height by 10pt
        if template_type in ["large", "large_eco", "large_nonaccredited"]:
            # Large template scope adjustment
            original_scope = coords["Scope"]
            adjusted_scope = fitz.Rect(
                original_scope.x0,      # x0: 85 (unchanged)
                original_scope.y0,      # y0: 351 (unchanged)  
                original_scope.x1,      # x1: 577 (unchanged)
                original_scope.y1 - 10  # y1: 536 - 10 = 526 (reduced by 10pt)
            )
            coords["Scope"] = adjusted_scope
            print(f"🔍 [CERTIFICATE] Large template scope adjusted: {original_scope} → {adjusted_scope}")
            
        elif template_type in ["standard", "standard_eco", "standard_nonaccredited"]:
            # Standard template scope adjustment (both short and long)
            original_short = coords["Scope"]["short"]
            original_long = coords["Scope"]["long"]
            
            adjusted_short = fitz.Rect(
                original_short.x0,      # x0: 87.9 (unchanged)
                original_short.y0,      # y0: 386 (unchanged)
                original_short.x1,      # x1: 580 (unchanged)
                original_short.y1 - 10  # y1: 475 - 10 = 465 (reduced by 10pt)
            )
            
            adjusted_long = fitz.Rect(
                original_long.x0,       # x0: 87.9 (unchanged)
                original_long.y0,       # y0: 373 (unchanged)
                original_long.x1,       # x1: 580 (unchanged)
                original_long.y1 - 10   # y1: 486 - 10 = 476 (reduced by 10pt)
            )
            
            coords["Scope"]["short"] = adjusted_short
            coords["Scope"]["long"] = adjusted_long
            print(f"🔍 [CERTIFICATE] Standard template scope adjusted: short={adjusted_short}, long={adjusted_long}")
            
        elif template_type in ["logo", "logo_nonaccredited", "logo_other", "logo_nonaccredited_other"]:
            # Logo template scope adjustment (same as standard)
            original_short = coords["Scope"]["short"]
            original_long = coords["Scope"]["long"]
            
            adjusted_short = fitz.Rect(
                original_short.x0,      # x0: 87.9 (unchanged)
                original_short.y0,      # y0: 426 (unchanged)
                original_short.x1,      # x1: 580 (unchanged)
                original_short.y1 - 10  # y1: 515 - 10 = 505 (reduced by 10pt)
            )
            
            adjusted_long = fitz.Rect(
                original_long.x0,       # x0: 87.9 (unchanged)
                original_long.y0,       # y0: 413 (unchanged)
                original_long.x1,       # x1: 580 (unchanged)
                original_long.y1 - 10   # y1: 526 - 10 = 516 (reduced by 10pt)
            )
            
            coords["Scope"]["short"] = adjusted_short
            coords["Scope"]["long"] = adjusted_long
            print(f"🔍 [CERTIFICATE] Logo template scope adjusted: short={adjusted_short}, long={adjusted_long}")
            
    else:
        print(f"🔍 [CERTIFICATE] No Extra Line - using standard scope coordinates")
    
    # Function to insert logo with smart positioning
    def insert_logo_with_smart_positioning(page, logo_image, logo_rect):
        """
        Smart logo insertion that handles different aspect ratios:
        - Square logos: Use full coordinates, maintain aspect ratio
        - Horizontal logos: Center horizontally, maintain aspect ratio
        - Vertical logos: Center vertically, maintain aspect ratio
        """
        # Get logo dimensions
        logo_width = logo_image.width
        logo_height = logo_image.height
        
        # Calculate aspect ratios
        logo_aspect = logo_width / logo_height
        rect_aspect = logo_rect.width / logo_rect.height
        
        # Determine positioning strategy based on aspect ratio
        if logo_aspect > rect_aspect:
            # Logo is more horizontal than rectangle
            # Scale to fit width, center vertically
            scale_factor = logo_rect.width / logo_width
            new_width = logo_rect.width
            new_height = logo_height * scale_factor
            
            # Center vertically
            x = logo_rect.x0
            y = logo_rect.y0 + (logo_rect.height - new_height) / 2
            
        elif logo_aspect < rect_aspect:
            # Logo is more vertical than rectangle
            # Scale to fit height, center horizontally
            scale_factor = logo_rect.height / logo_height
            new_width = logo_width * scale_factor
            new_height = logo_rect.height
            
            # Center horizontally
            x = logo_rect.x0 + (logo_rect.width - new_width) / 2
            y = logo_rect.y0
            
        else:
            # Logo and rectangle have similar aspect ratios
            # Scale to fit within rectangle, center both ways
            scale_factor = min(logo_rect.width / logo_width, logo_rect.height / logo_height)
            new_width = logo_width * scale_factor
            new_height = logo_height * scale_factor
            
            # Center both horizontally and vertically
            x = logo_rect.x0 + (logo_rect.width - new_width) / 2
            y = logo_rect.y0 + (logo_rect.height - new_height) / 2
        
        # Convert PIL Image to PyMuPDF image
        try:
            # Convert PIL Image to bytes
            import io
            img_buffer = io.BytesIO()
            logo_image.save(img_buffer, format='PNG')
            img_buffer.seek(0)
            img_bytes = img_buffer.read()
            
            # Insert into PDF
            page.insert_image(fitz.Rect(x, y, x + new_width, y + new_height), stream=img_bytes)
            print(f"🔍 [LOGO] Logo inserted with smart positioning: {new_width:.1f}x{new_height:.1f}")
            
        except Exception as e:
            print(f"❌ [LOGO] Error inserting logo: {e}")

    # ✅ ADDED: Shared logo functions for better logo handling
    def insert_logo_into_pdf(page, logo_file, logo_rect):
        """
        Insert logo into PDF with smart positioning
        """
        try:
            # Convert logo file to image
            logo_image = convert_file_to_image(logo_file)
            # Use smart positioning logic
            insert_logo_with_smart_positioning(page, logo_image, logo_rect)
            print(f"✅ [LOGO] Logo inserted successfully: {logo_file.name if hasattr(logo_file, 'name') else 'unknown'}")
        except Exception as e:
            print(f"❌ [LOGO] Failed to insert logo: {e}")

    def convert_file_to_image(file):
        """
        Convert uploaded file to PIL Image
        """
        try:
            from PIL import Image
            import io
            
            if hasattr(file, 'file'):
                # Reset file pointer
                file.file.seek(0)
                # Read file content
                file_content = file.file.read()
                # Convert to PIL Image
                logo_image = Image.open(io.BytesIO(file_content))
                return logo_image
            else:
                raise ValueError("File object has no file attribute")
        except Exception as e:
            print(f"❌ [LOGO] Error converting file to image: {e}")
            raise

    # ✅ ADDED: Render optional fields function
    def render_optional_fields(page, values, key_coords, value_coords, font_settings):
        """
        Render optional fields with dynamic positioning based on available data.
        """
        # ✅ ADDED: Defensive checks for None values
        if values is None:
            print(f"⚠️ [CERTIFICATE] Values dictionary is None in render_optional_fields - skipping")
            return
        
        if key_coords is None or value_coords is None:
            print(f"⚠️ [CERTIFICATE] Coordinates are None in render_optional_fields - skipping")
            return
        
        if font_settings is None:
            print(f"⚠️ [CERTIFICATE] Font settings is None in render_optional_fields - skipping")
            return
        # Define field order (top to bottom) and their display labels
        fields = [
            "Certificate Number",        # seq 1 - TOP (always present)
            "Initial Registration Date", # seq 2 (optional - not always present)
            "Original Issue Date",       # seq 3 (always present)
            "Issue Date",               # seq 4 (always present)
            "Surveillance Group",       # seq 5 (only 1 of 3 fields present)
            "Recertification Date"      # seq 6 - BOTTOM (always present)
        ]
        
        # Define the surveillance group fields (only 1 will be present)
        surveillance_group_fields = [
            "Surveillance/ Expiry Date",
            "Surveillance Due Date", 
            "Expiry Date"
        ]
        
        # Custom display labels for PDF rendering (cleaner, shorter)
        display_labels = {
            "Certificate Number": "Certificate No.",
            "Initial Registration Date": "Initial Registration Date",
            "Original Issue Date": "Original Issue Date",
            "Issue Date": "Issue Date", 
            "Surveillance/ Expiry Date": "Surveillance/ Expiry Date",
            "Surveillance Due Date": "Surveillance Due Date",
            "Expiry Date": "Expiry Date",
            "Recertification Date": "Recertification Date"
        }

        # Filter available fields (non-empty) with special handling for surveillance group
        available_fields = []
        
        # ✅ ADDED: Debug logging for values received
        print(f"🔍 [CERTIFICATE] ===== OPTIONAL FIELDS DEBUG =====")
        print(f"🔍 [CERTIFICATE] Values received: {list(values.keys())}")
        print(f"🔍 [CERTIFICATE] Surveillance/ Expiry Date: '{values.get('Surveillance/ Expiry Date', '')}'")
        print(f"🔍 [CERTIFICATE] Surveillance Due Date: '{values.get('Surveillance Due Date', '')}'")
        print(f"🔍 [CERTIFICATE] Expiry Date: '{values.get('Expiry Date', '')}'")
        
        for field in fields:
            if field == "Surveillance Group":
                # Handle surveillance group - find which field is present
                surveillance_value = None
                surveillance_label = None
                
                print(f"🔍 [CERTIFICATE] Processing Surveillance Group...")
                for surveillance_field in surveillance_group_fields:
                    print(f"🔍 [CERTIFICATE] Checking '{surveillance_field}': '{values.get(surveillance_field, '')}'")
                    if surveillance_field in values and values[surveillance_field]:
                        surveillance_value = values[surveillance_field]
                        surveillance_label = surveillance_field
                        print(f"🔍 [CERTIFICATE] Found surveillance field: '{surveillance_label}' = '{surveillance_value}'")
                        break
                
                if surveillance_value and surveillance_label:
                    available_fields.append((surveillance_label, surveillance_value))
                    print(f"🔍 [CERTIFICATE] Added surveillance field to available fields")
                else:
                    print(f"🔍 [CERTIFICATE] No surveillance field found or all are empty")
            else:
                value = values.get(field, "").strip()
                if value:  # Only include non-empty fields
                    available_fields.append((field, value))
                    print(f"🔍 [CERTIFICATE] Added field '{field}' = '{value}' to available fields")
                else:
                    print(f"🔍 [CERTIFICATE] Field '{field}' is empty, skipping")
        
        print(f"🔍 [CERTIFICATE] Total available fields: {len(available_fields)}")
        print(f"🔍 [CERTIFICATE] Available fields: {available_fields}")
        print(f"🔍 [CERTIFICATE] ===== END OPTIONAL FIELDS DEBUG =====")

        # Calculate starting position using formula: (6 - available_count) + 1
        total_fields = 6
        available_count = len(available_fields)
        starting_position = (6 - available_count) + 1

        # Essential logging only
        print(f"🔍 [CERTIFICATE] Optional fields: {available_count}/{total_fields} available")

        # Render from starting position
        for i, (field, value) in enumerate(available_fields):
            coord_index = starting_position - 1 + i  # Convert to 0-based index

            if coord_index < len(key_coords):  # Prevent overflow
                # Prepare text with custom display labels
                key_text = display_labels[field]  # Use custom display label
                value_text = f":{value}"            # Colon + value (no space)

                # Essential field logging
                print(f"🔍 [CERTIFICATE] Rendering: {field}")

                # Extract (x, y) coordinates from rectangles
                key_x = key_coords[coord_index].x0
                key_y = key_coords[coord_index].y0
                value_x = value_coords[coord_index].x0
                value_y = value_coords[coord_index].y0

                # ✅ ADDED: Defensive checks for font settings
                fontsize = font_settings.get('fontsize', 15)  # Default to 15pt
                fontname = font_settings.get('fontname', 'Times-Roman')  # Default to Times-Roman
                color = font_settings.get('color', (0, 0, 0))  # Default to black
                
                # Insert at respective coordinates using (x, y) points
                page.insert_text(
                    (key_x, key_y),
                    key_text,
                    fontsize=fontsize,
                    fontname=fontname,
                    color=color
                )

                page.insert_text(
                    (value_x, value_y),
                    value_text,
                    fontsize=fontsize,
                    fontname=fontname,
                    color=color
                )

                print(f"   ✅ Rendered successfully at position {coord_index + 1}")
            else:
                print(f"⚠️ [CERTIFICATE] Warning: Coordinate index {coord_index} out of bounds")

        print(f"🔍 [CERTIFICATE] ===== END OPTIONAL FIELDS ANALYSIS =====\n")

    # ✅ ADDED: Optional field coordinates for large template
    large_optional_key_coordinates = [
        fitz.Rect(175.5, 522, 343, 530),    # Row 1: Certificate Number
        fitz.Rect(175.5, 538, 343, 548),    # Row 2: Initial Registration Date
        fitz.Rect(175.5, 556, 343, 566),    # Row 3: Original Issue Date
        fitz.Rect(175.5, 574, 343, 584),    # Row 4: Issue Date
        fitz.Rect(175.5, 592, 343, 602),    # Row 5: Surveillance Group (only 1 field)
        fitz.Rect(175.5, 610, 343, 620)     # Row 6: Recertification Date
    ]
    
    large_optional_value_coordinates = [
        fitz.Rect(362.1, 522, 446.4, 530),  # Row 1: Certificate Number value
        fitz.Rect(362.1, 538, 446.4, 548),  # Row 2: Initial Registration Date value
        fitz.Rect(362.1, 556, 446.4, 566),  # Row 3: Original Issue Date value
        fitz.Rect(362.1, 574, 446.4, 584),  # Row 4: Issue Date value
        fitz.Rect(362.1, 592, 446.4, 602),  # Row 5: Surveillance Group value (only 1 field)
        fitz.Rect(362.1, 610, 446.4, 620)   # Row 6: Recertification Date value
    ]
    
    # ✅ ADDED: Optional field coordinates for standard template
    standard_optional_key_coordinates = [
        fitz.Rect(175.5, 499.1, 343, 509.1),    # Row 1: Certificate Number
        fitz.Rect(175.5, 516.9, 343, 526.9),    # Row 2: Initial Registration Date
        fitz.Rect(175.5, 535.1, 343, 545.1),    # Row 3: Original Issue Date
        fitz.Rect(175.5, 553.9, 343, 563.9),    # Row 4: Issue Date
        fitz.Rect(175.5, 571.6, 343, 581.6),    # Row 5: Surveillance Group (only 1 field)
        fitz.Rect(175.5, 589.3, 343, 599.3)     # Row 6: Recertification Date
    ]
    
    standard_optional_value_coordinates = [
        fitz.Rect(362.1, 499.1, 446.4, 509.1),  # Row 1: Certificate Number value
        fitz.Rect(362.1, 516.9, 446.4, 526.9),  # Row 2: Initial Registration Date value
        fitz.Rect(362.1, 535.1, 446.4, 545.1),  # Row 3: Original Issue Date value
        fitz.Rect(362.1, 553.9, 446.4, 563.9),  # Row 4: Issue Date value
        fitz.Rect(362.1, 571.6, 446.4, 581.6),  # Row 5: Surveillance Group value (only 1 field)
        fitz.Rect(362.1, 589.3, 446.4, 599.3)   # Row 6: Recertification Date value
    ]
    
    # ✅ UPDATED: Font settings for optional fields (matching soft copy)
    optional_font_settings = {
        "fontname": "Times-Roman",  # Same as soft copy
        "fontsize": 13,             # Same as soft copy (was 15pt)
        "color": (0, 0, 0)         # Black
    }

    # ✅ ADDED: Adjust scope coordinates based on whether Initial Registration Date is present
    initial_registration_date = values.get("Initial Registration Date", "")
    if initial_registration_date and initial_registration_date.strip():
        # When Initial Registration Date is present, reduce scope height to accommodate the extra field
        print(f"🔍 [CERTIFICATE] Initial Registration Date present - adjusting scope coordinates for large template")
        if template_type in ["large", "large_eco", "large_nonaccredited"]:
            # Adjust large template scope coordinates
            coords["Scope"] = fitz.Rect(85, 351, 577, 520)  # Reduced height by 16 units
            print(f"🔍 [CERTIFICATE] Adjusted large template scope coordinates for Initial Registration Date")
    else:
        print(f"🔍 [CERTIFICATE] Using standard scope coordinates (Initial Registration Date not present)")

    # Process optional fields
    if template_type in ["large", "large_eco", "large_nonaccredited"]:
        render_optional_fields(page, values, large_optional_key_coordinates, large_optional_value_coordinates, optional_font_settings)
    else:
        render_optional_fields(page, values, standard_optional_key_coordinates, standard_optional_value_coordinates, optional_font_settings)

    # Calculate optional fields count for field processing
    optional_fields_count = 0
    # ✅ ADDED: Defensive check for values before optional fields calculation
    if values is not None:
        for field in ["Certificate Number", "Initial Registration Date", "Original Issue Date", "Issue Date", "Surveillance/ Expiry Date", "Surveillance Due Date", "Expiry Date", "Recertification Date"]:
            if values.get(field, "").strip():
                optional_fields_count += 1
    else:
        print(f"⚠️ [CERTIFICATE] Values is None during optional fields calculation - using 0")
    
    # Process each field with enhanced text handling
    # ✅ FIELD CLASSIFICATION SYSTEM:
    # ==========================================
    # RENDERING_FIELDS: Main content fields rendered using main coordinates
    # OPTIONAL_FIELDS: Additional fields rendered using separate optional coordinates  
    # METADATA_FIELDS: Template selection fields (NOT rendered on PDF)
    # ==========================================
    RENDERING_FIELDS = ["Company Name", "Scope", "ISO Standard", "management_system", "Extra Line"]
    OPTIONAL_FIELDS = ["Certificate Number", "Initial Registration Date", "Original Issue Date", "Issue Date", "Surveillance/ Expiry Date", "Surveillance Due Date", "Expiry Date", "Recertification Date"]
    METADATA_FIELDS = ["Size", "Accreditation", "Logo", "Country", "logo_lookup"]
    
    # ✅ FIELD PROCESSING FLOW:
    # 1. Metadata fields (Size, Accreditation, Logo, Country) -> SKIPPED (no rendering needed)
    # 2. Optional fields (Certificate Number, Dates, etc.) -> SKIPPED (rendered via render_optional_fields)
    # 3. Main rendering fields (Company Name, Scope, ISO Standard, management_system) -> PROCESSED (main coordinates)
    
    # ✅ ADDED: Defensive check for values before main field processing
    if values is None:
        print(f"⚠️ [CERTIFICATE] Values dictionary is None in main field processing - cannot continue")
        raise ValueError("Values dictionary is None - cannot generate certificate")
    
    # Scope text now uses justification (left and right alignment) for professional appearance
    for field, text in values.items():
        # ✅ ADDED: Skip metadata fields that don't need rendering
        if field in METADATA_FIELDS:
            continue
            
        # ✅ ADDED: Skip optional fields as they are handled separately
        if field in OPTIONAL_FIELDS:
            continue
        
        if field == "Company Name":
            company_text = str(text) if text is not None else ""
            
            # Handle Company Name and Address together
            address_text = values.get("Address", "")
            safe_address_text = str(address_text) if address_text is not None else ""
            


            
            # PRE-PROCESS: Apply line break logic BEFORE font size calculation
            # This ensures both font calculation and rendering use the same processed text
            
            # Process Company Name with line break preservation
            def process_text_with_line_breaks(input_text, field_name):
                if not input_text:
                    return []
                
                # ✅ UPDATED: Split by actual line breaks and PRESERVE empty lines
                lines = input_text.split('\n')
                processed_lines = []
                
                for line in lines:
                    # Preserve empty lines to maintain spacing from Excel
                    processed_lines.append(line)  # Keep original line (including empty ones)

                
                return processed_lines
            
            # Pre-process both company and address text to handle line breaks
            company_processed_lines = process_text_with_line_breaks(company_text, "Company")
            address_processed_lines = process_text_with_line_breaks(safe_address_text, "Address")
            
            # ✅ ADDED: Defensive check for coords dictionary
            if coords is None:
                print(f"⚠️ [CERTIFICATE] Coords dictionary is None - cannot continue")
                raise ValueError("Coords dictionary is None - cannot generate certificate")
            
            if "Company Name and Address" not in coords:
                print(f"⚠️ [CERTIFICATE] Company Name and Address coordinates not found - cannot continue")
                raise ValueError("Company Name and Address coordinates not found - cannot generate certificate")
            
            rect = coords["Company Name and Address"]
            
            # Check if address text is empty or None
            if not safe_address_text:
                print(f"[WARNING] [COMPANY ADDRESS] WARNING: Address text is empty or None!")
            elif safe_address_text.strip() == "":
                print(f"[WARNING] [COMPANY ADDRESS] WARNING: Address text is only whitespace!")
            else:
                print(f"[SUCCESS] [COMPANY ADDRESS] Address text is valid and non-empty")
            
            # Combine text with natural spacing (single line break)
            combined_text = f"{company_text}\n{safe_address_text}"  # \n creates ~2-3pt spacing
            
            # ✅ ADDED: Defensive check for font_starts dictionary
            if font_starts is None:
                print(f"⚠️ [CERTIFICATE] Font_starts dictionary is None - using default")
                start_size = 30
            else:
                start_size = font_starts.get("Company Name and Address", 30)
            font_size = start_size
            
    
            

            
            # ✅ UPDATED: Dynamic Company Name font sizing based on line count
            # First, determine if Company Name will be single line or multi-line
            company_lines_count = len([line for line in company_processed_lines if line.strip()])
            
            # Set initial font size based on line count
            if company_lines_count <= 1:
                company_font_size = 35  # Single line - start with 35pt
                print(f"🔍 [CERTIFICATE] Company Name: Single line detected, starting with {company_font_size}pt")
            else:
                company_font_size = 30  # Multiple lines - start with 30pt
                print(f"🔍 [CERTIFICATE] Company Name: {company_lines_count} lines detected, starting with {company_font_size}pt")
            
            address_font_size = 13.6
            
            # Variables to store the final wrapped lines and font sizes
            final_company_lines = []
            final_address_lines = []
            
            # ✅ IMPROVED: Different logic for single line vs multi-line company names
            if company_lines_count <= 1:
                # NO cmd+enter in Excel: Force single line, use font reduction only
                print(f"🔍 [CERTIFICATE] No cmd+enter detected - forcing single line with font reduction")
                
                while company_font_size >= 8:  # Minimum font size
                    # Check if entire company name fits in one line at current font size
                    font_obj = fitz.Font(fontname=fontname)
                    text_width = font_obj.text_length(company_text, company_font_size)
                    
                    if text_width <= rect.width - 10:  # Leave margin
                        # Text fits in one line - use this font size
                        final_company_lines = [company_text]  # Single line
                        print(f"✅ [CERTIFICATE] Company name fits in one line at {company_font_size}pt (width: {text_width:.1f}pt)")
                        break
                    else:
                        # Text too wide - reduce font size and try again
                        print(f"🔍 [CERTIFICATE] Company name too wide at {company_font_size}pt (width: {text_width:.1f}pt > {rect.width - 10:.1f}pt), reducing to {company_font_size - 1}pt")
                        company_font_size -= 1
                
                # If we reached minimum font size and still doesn't fit, use the minimum
                if company_font_size < 8:
                    company_font_size = 8
                    final_company_lines = [company_text]
                    print(f"⚠️ [CERTIFICATE] Company name forced to minimum font size 8pt")
                
            else:
                # cmd+enter present in Excel: Allow word wrapping up to 2 lines
                print(f"🔍 [CERTIFICATE] cmd+enter detected - allowing word wrapping up to 2 lines")
                
            while company_font_size >= 8:  # Minimum font size
                company_lines = []
                
                # Process each pre-processed line with word wrapping
                for processed_line in company_processed_lines:
                    if not processed_line.strip():  # Empty line - preserve it
                        company_lines.append("")  # Add empty line to maintain spacing
                        continue
                    
                    # Non-empty line - apply word wrapping
                    words = processed_line.split()
                    current_line = ""
                    
                    for word in words:
                        test_line = current_line + (" " if current_line else "") + word
                        font_obj = fitz.Font(fontname=fontname)
                        if font_obj.text_length(test_line, company_font_size) <= rect.width - 10:  # Leave margin
                            current_line = test_line
                        else:
                            if current_line:
                                company_lines.append(current_line)
                            current_line = word
                    
                    if current_line:
                        company_lines.append(current_line)
                
                # ✅ UPDATED: Allow Company Name to use up to 2 lines (after line breaks + word wrapping)
                if len(company_lines) <= 2:
                    final_company_lines = company_lines.copy()
                    break
                
                # Reduce Company Name font size
                company_font_size -= 1
            
            # Calculate Company Name height
            # Template-specific line spacing: 1.1 for large/logo, 1.2 for standard
            if template_type in ["large", "large_eco", "large_nonaccredited", "logo", "logo_nonaccredited", "logo_other", "logo_other_nonaccredited"]:
                company_height = len(final_company_lines) * company_font_size * 1.1  # Tight spacing for large/logo templates
            else:  # standard templates
                company_height = len(final_company_lines) * company_font_size * 1.2  # Loose spacing for standard templates
            
            # Now find font size for Address to fit in remaining space
            remaining_height = rect.height - company_height - 2  # Leave margin

            
            address_font_size_attempts = 0
            while address_font_size >= 6:  # Minimum font size
                address_font_size_attempts += 1

                
                # Process Address using pre-processed lines with word wrapping
                address_lines = []
                
                # Process each pre-processed address line with word wrapping
                for processed_line in address_processed_lines:
                    if not processed_line.strip():  # Empty line - preserve it
                        address_lines.append("")  # Add empty line to maintain spacing
                        continue
                    
                    # Non-empty line - apply word wrapping
                    words = processed_line.split()
                    current_line = ""
                    
                    for word in words:
                        test_line = current_line + (" " if current_line else "") + word
                        font_obj = fitz.Font(fontname=fontname)
                        test_width = font_obj.text_length(test_line, address_font_size)
                        if test_width <= rect.width - 10:  # Leave margin
                            current_line = test_line
                        else:
                            if current_line:
                                address_lines.append(current_line)
                            current_line = word
                    
                    if current_line:
                        address_lines.append(current_line)
                
                # Calculate Address height
                # Template-specific line spacing: 1.1 for large/logo, 1.2 for standard
                if template_type in ["large", "large_eco", "large_nonaccredited", "logo", "logo_nonaccredited", "logo_other", "logo_other_nonaccredited"]:
                    address_height = len(address_lines) * address_font_size * 1.1  # Tight spacing for large/logo templates
                else:  # standard templates
                    address_height = len(address_lines) * address_font_size * 1.2  # Loose spacing for standard templates

                
                # Check if Address fits in remaining space
                if address_height <= remaining_height:
                    final_address_lines = address_lines.copy()
                    print(f"[SUCCESS] [COMPANY ADDRESS] Address fits! Final font size: {address_font_size}pt")
                    break
                else:
                    print(f"[ERROR] [COMPANY ADDRESS] Address too tall: {address_height:.1f}pt > {remaining_height:.1f}pt, reducing font size")
                
                # Reduce Address font size
                address_font_size -= 0.5
            
            # Now render Company Name and Address dynamically
            if final_company_lines or final_address_lines:
                


                
                # Calculate total height
                total_height = company_height + address_height

                
                # Fixed top margin always 2pt
                start_y = rect.y0 + 2  # 2pt margin from top of box

                
                # Render Company Name first (starts from top)
                current_y = start_y
                for i, line in enumerate(final_company_lines):
                    # Template-specific line spacing: 1.1 for large/logo, 1.2 for standard
                    if template_type in ["large", "large_eco", "large_nonaccredited", "logo", "logo_nonaccredited", "logo_other", "logo_other_nonaccredited"]:
                        line_height = company_font_size * 1.1  # Tight spacing for large/logo templates
                    else:  # standard templates
                        line_height = company_font_size * 1.2  # Loose spacing for standard templates
                    y_pos = current_y + line_height/2  # Adjust for baseline
                    center_x = (rect.x0 + rect.x1) / 2
                    
                                        # ✅ UPDATED: Handle empty lines (preserve spacing from Excel)
                    if line.strip():  # Non-empty line - render text
                        # ✅ ENHANCED: Use mixed format text rendering for bold detection
                        if '**' in line or '__' in line:
                            # Calculate total width for centering
                            segments = process_bold_text(line)
                            total_width = 0
                            for segment_text, _, _ in segments:
                                if segment_text:
                                    font_obj = fitz.Font(fontname="Times-Bold" if "**" in segment_text or "__" in segment_text else "Times-Roman")
                                    total_width += font_obj.text_length(segment_text, company_font_size)
                            
                            x_pos = center_x - total_width / 2
                            render_mixed_format_text(page, (x_pos, y_pos), line, company_font_size, color)
                        else:
                            # Standard rendering for non-bold text
                            line_width = font_obj.text_length(line, company_font_size)
                            x_pos = center_x - line_width / 2
                            
                            # ✅ ADDED: Dynamic font selection for Company Name
                            dynamic_font = get_font_for_text(line, fontname)
                            page.insert_text(
                                (x_pos, y_pos),
                                line,
                                fontsize=company_font_size,
                                fontname=dynamic_font,
                                color=color
                            )
                    # Empty line - just advance position (creates blank space)
                    
                    current_y += line_height
                
                # Add spacing between Company Name and Address
                spacing = 8  # 8pt spacing between Company Name and Address
                current_y += spacing
                
                # Now render Address below Company Name (fills remaining space)
                for i, line in enumerate(final_address_lines):
                    # Template-specific line spacing: 1.1 for large/logo, 1.2 for standard
                    if template_type in ["large", "large_eco", "large_nonaccredited", "logo", "logo_nonaccredited", "logo_other", "logo_other_nonaccredited"]:
                        line_height = address_font_size * 1.1  # Tight spacing for large/logo templates
                    else:  # standard templates
                        line_height = address_font_size * 1.2  # Loose spacing for standard templates
                    y_pos = current_y + line_height/2  # Adjust for baseline
                    center_x = (rect.x0 + rect.x1) / 2
                    
                    # ✅ UPDATED: Handle empty lines (preserve spacing from Excel)
                    if line.strip():  # Non-empty line - render text
                        # ✅ ENHANCED: Use mixed format text rendering for bold detection
                        if '**' in line or '__' in line:
                            # Calculate total width for centering
                            segments = process_bold_text(line)
                            total_width = 0
                            for segment_text, _, _ in segments:
                                if segment_text:
                                    font_obj = fitz.Font(fontname="Times-Bold" if "**" in segment_text or "__" in segment_text else "Times-Roman")
                                    total_width += font_obj.text_length(segment_text, address_font_size)
                            
                            x_pos = center_x - total_width / 2
                            render_mixed_format_text(page, (x_pos, y_pos), line, address_font_size, color)
                        else:
                            # Standard rendering for non-bold text
                            line_width = font_obj.text_length(line, address_font_size)
                            x_pos = center_x - line_width / 2
                            
                            # ✅ ADDED: Dynamic font selection for Address
                            dynamic_font = get_font_for_text(line, fontname)
                            page.insert_text(
                                (x_pos, y_pos),
                                line,
                                fontsize=address_font_size,
                                fontname=dynamic_font,
                                color=color
                            )
                    # Empty line - just advance position (creates blank space)
                    
                    current_y += line_height
                
                # Print final font sizes for Company Name and Address

                
                # Skip Address processing since it's handled above
                continue
            
        elif field == "Address":
            # Skip Address processing since it's handled with Company Name
            continue
        
        elif field == "ISO Standard":
            # After processing ISO Standard, render the management system line
            iso_standard_text = text
            
            # Expand the ISO standard to full version with year
            expanded_iso = expand_iso_standard(iso_standard_text)
            
            # Get the description from the mapping using the expanded version
            system_name = ISO_STANDARDS_DESCRIPTIONS.get(expanded_iso, "Management System")
            
            # Capitalize first letters of each word in system_name
            system_name_caps = ' '.join(word.capitalize() for word in system_name.split())
            
            # Create the management system line
            management_line = f"This is to certify that the {system_name_caps} of"
            
            # Get the management_system rectangle
            management_rect = coords["management_system"]
            

            
            # Calculate center position for the text
            center_x = (management_rect.x0 + management_rect.x1) / 2
            center_y = (management_rect.y0 + management_rect.y1) / 2 + 15/3  # Adjust for baseline
            
            # Calculate text width for centering
            font_obj = fitz.Font(fontname="Times-BoldItalic")  # Use bold italic font
            text_width = font_obj.text_length(management_line, 15)
            start_x = center_x - text_width / 2
            

            
            # Insert the management system text
            page.insert_text(
                (start_x, center_y),
                management_line,
                fontsize=15,
                fontname="Times-BoldItalic", # Bold italic font
                color=(0, 0, 0)  # Black color
            )
            
            # Print font size for management system

            
            # Update the text to use expanded version for display
            text = expanded_iso
        
        # ✅ VALIDATION: Ensure only rendering fields are processed
        if field not in RENDERING_FIELDS:
            continue
        
        # Handle other fields normally
        if field not in coords:
            continue
            
        rect = coords[field]
        
        # Template-specific starting font size for Scope
        if field == "Scope" and template_type == "standard" and scope_layout == "short":
            start_size = 15  # Standard template short scope: max 15pt
        else:
            # ✅ ADDED: Defensive check for font_starts dictionary
            if font_starts is None:
                print(f"⚠️ [CERTIFICATE] Font_starts dictionary is None - using default for field '{field}'")
                start_size = 30
            else:
                start_size = font_starts.get(field, 30)  # Use existing logic for other cases
        
        font_size = start_size
        

        

        
        # Reduce font size if it doesn't fit, but ensure minimum size
        while font_size >= 12:  # Increased minimum from 10 to 12
            text_height = get_text_height(text, font_size, fontname, rect.width)
            limit = rect.height if field != "Company Name" else rect.height * 2

            if text_height <= limit:
                break
            font_size -= 1
        

        
        if field == "Scope":
            # PowerPoint-style centering with automatic font size reduction
            original_font_size = font_size
            
            # 🔍 [CERTIFICATE DEBUG] Scope processing start
            print(f"🔍 [CERTIFICATE DEBUG] ===== SCOPE ANALYSIS START =====")
            print(f"🔍 [CERTIFICATE DEBUG] Template type: {template_type}")
            print(f"🔍 [CERTIFICATE DEBUG] Scope coordinates: {rect}")
            print(f"🔍 [CERTIFICATE DEBUG] Scope text length: {len(text)} characters")
            print(f"🔍 [CERTIFICATE DEBUG] Scope text preview: '{text[:100]}{'...' if len(text) > 100 else ''}'")
            
            # Reduce font size until text fits within box boundaries
            font_size = original_font_size
            iteration_count = 0
            min_font_size = 4  # Allow font size to go below 8pt if needed

            while font_size >= min_font_size:  # Changed from 8 to 4
                iteration_count += 1
                
                print(f"🔍 [CERTIFICATE DEBUG] Font size attempt {iteration_count}: {font_size}pt")
                
                # Enhanced text processing with bullet point detection AND line break preservation
                lines = []
                
                # Split text by actual line breaks first
                if '\n' in text:
                    # Handle line breaks properly
                    text_lines = text.split('\n')
                    for text_line in text_lines:
                        if text_line.strip():  # Only process non-empty lines
                            # Process each line for word wrapping and bullet points
                            words = text_line.strip().split()
                            current_line = ""
                            
                            for word in words:
                                # Check if word starts with bullet point indicators
                                is_bullet_point = any(word.startswith(indicator) for indicator in ['-', '•', '>', '→', '▪', '▫', '*'])
                                
                                # If it's a bullet point and we have content, start a new line
                                if is_bullet_point and current_line:
                                    lines.append(current_line)
                                    # Replace * with • for display (in font size calculation)
                                    if word.startswith('*'):
                                        safe_word = str(word) if word is not None else ""
                                        print(f"[BULLET] Will replace '{safe_word}' with '•{safe_word[1:]}' in final rendering")
                                    current_line = word
                                    continue
                                
                                test_line = current_line + (" " if current_line else "") + word
                                font_obj = fitz.Font(fontname=fontname)
                                
                                if font_obj.text_length(test_line, font_size) <= rect.width:
                                    current_line = test_line
                                else:
                                    if current_line:
                                        lines.append(current_line)
                                    current_line = word
                            
                            if current_line:
                                lines.append(current_line)
                else:
                    # No line breaks, process as single block
                    words = text.split()
                    current_line = ""
                    
                    for word in words:
                        # Check if word starts with bullet point indicators
                        is_bullet_point = any(word.startswith(indicator) for indicator in ['-', '•', '>', '→', '▪', '▫', '*'])
                        
                        # If it's a bullet point and we have content, start a new line
                        if is_bullet_point and current_line:
                            lines.append(current_line)
                            # Replace * with • for display (in font size calculation)
                            if word.startswith('*'):
                                safe_word = str(word) if word is not None else ""
                                print(f"[BULLET] Will replace '{safe_word}' with '•{safe_word[1:]}' in final rendering")
                            current_line = word
                            continue
                        
                        test_line = current_line + (" " if current_line else "") + word
                        font_obj = fitz.Font(fontname=fontname)
                        
                        if font_obj.text_length(test_line, font_size) <= rect.width:
                            current_line = test_line
                        else:
                            if current_line:
                                lines.append(current_line)
                            current_line = word
                    
                    if current_line:
                        lines.append(current_line)
                
                # Calculate total height of all lines
                # Template-specific line spacing: 1.1 for large/logo, 1.2 for standard
                if template_type in ["large", "large_eco", "large_nonaccredited", "logo", "logo_nonaccredited", "logo_other", "logo_other_nonaccredited"]:
                    line_height = font_size * 1.1  # Tight spacing for large/logo templates
                else:  # standard templates
                    line_height = font_size * 1.2  # Loose spacing for standard templates
                total_height = len(lines) * line_height
                
                print(f"🔍 [CERTIFICATE DEBUG] Calculated height: {total_height:.1f}pt (lines: {len(lines)}, line_height: {line_height:.1f}pt)")
                print(f"🔍 [CERTIFICATE DEBUG] Available height: {rect.height:.1f}pt")
                print(f"🔍 [CERTIFICATE DEBUG] Height utilization: {(total_height/rect.height)*100:.1f}%")
                
                # Check if text fits vertically within box boundaries
                if total_height <= rect.height:  # No margin
                    print(f"🔍 [CERTIFICATE DEBUG] ✅ Text fits! Using font size: {font_size}pt")
                    break
                else:
                    print(f"🔍 [CERTIFICATE DEBUG] ❌ Text overflow: {total_height:.1f}pt > {rect.height:.1f}pt, reducing font size")
                
                font_size -= 1

            # Check if we hit the minimum font size and still have overflow
            if font_size == min_font_size and total_height > rect.height:
                # Calculate how much overflow we have
                overflow_amount = total_height - rect.height
                overflow_percentage = (overflow_amount / rect.height) * 100
                
                # Add to overflow warnings
                company_name = values.get("Company Name", "Unknown Company")
                iso_standard = values.get("ISO Standard", "Unknown Standard")
                
                warning_msg = f"[OVERFLOW] {company_name} - {iso_standard}: Scope text exceeds coordinates by {overflow_percentage:.1f}% (font size reduced to {min_font_size}pt)"
                overflow_warnings.append({
                    "company_name": company_name,
                    "iso_standard": iso_standard,
                    "overflow_percentage": overflow_percentage,
                    "final_font_size": min_font_size,
                    "message": warning_msg
                })
                
                print(warning_msg)
                print(f"[OVERFLOW] Text will be truncated to fit within coordinates")
                
                # Force text to fit by truncating if necessary
                max_lines = int(rect.height / min_font_size)
                if len(lines) > max_lines:
                    lines = lines[:max_lines]
                    print(f"[OVERFLOW] Truncated to {max_lines} lines to fit coordinates")

            # Now draw the text with the fitting font size using enhanced processing
            # Process text to handle line breaks and bullet points properly
            
            # Replace all asterisks with bullet points for display
            text = text.replace('*', '•')
            
            lines = []
            
            # Split text by actual line breaks first
            if '\n' in text:
                # Handle line breaks properly
                text_lines = text.split('\n')
                for text_line in text_lines:
                    if text_line.strip():  # Only process non-empty lines
                        # Process each line for word wrapping and bullet points
                        words = text_line.strip().split()
                        current_line = ""
                        
                        for word in words:
                            # Check if word starts with bullet point indicators
                            is_bullet_point = any(word.startswith(indicator) for indicator in ['-', '•', '>', '→', '▪', '▫', '*'])
                            
                            # If it's a bullet point and we have content, start a new line
                            if is_bullet_point and current_line:
                                lines.append(current_line)
                                # Replace * with • for display
                                display_word = word.replace('*', '•')
                                if word != display_word:
                                    print(f"[BULLET] Replaced '{word}' with '{display_word}'")
                                current_line = display_word
                                continue
                            
                            test_line = current_line + (" " if current_line else "") + word
                            font_obj = fitz.Font(fontname=fontname)
                            if font_obj.text_length(test_line, font_size) <= rect.width:
                                current_line = test_line
                            else:
                                if current_line:
                                    lines.append(current_line)
                                current_line = word
                        
                        if current_line:
                            lines.append(current_line)
            else:
                # No line breaks, process as single block
                words = text.split()
                current_line = ""
                
                for word in words:
                    # Check if word starts with bullet point indicators
                    is_bullet_point = any(word.startswith(indicator) for indicator in ['-', '•', '>', '→', '▪', '▫', '*'])
                    
                    # If it's a bullet point and we have content, start a new line
                    if is_bullet_point and current_line:
                        lines.append(current_line)
                        # Replace * with • for display
                        display_word = word.replace('*', '•')
                        current_line = display_word
                        continue
                    
                    test_line = current_line + (" " if current_line else "") + word
                    font_obj = fitz.Font(fontname=fontname)
                    if font_obj.text_length(test_line, font_size) <= rect.width:
                        current_line = test_line
                    else:
                        if current_line:
                            lines.append(current_line)
                        current_line = word
                
                if current_line:
                    lines.append(current_line)
            


            
            # Calculate total height and position vertically based on template type
            # Template-specific line spacing: 1.1 for large/logo, 1.2 for standard
            if template_type in ["large", "large_eco", "large_nonaccredited", "logo", "logo_nonaccredited", "logo_other", "logo_other_nonaccredited"]:
                line_height = font_size * 1.1  # Tight spacing for large/logo templates
            else:  # standard templates
                line_height = font_size * 1.2  # Loose spacing for standard templates
            total_height = len(lines) * line_height
            
            print(f"🔍 [CERTIFICATE DEBUG] ===== POSITIONING ANALYSIS =====")
            print(f"🔍 [CERTIFICATE DEBUG] Final font size: {font_size}pt")
            print(f"🔍 [CERTIFICATE DEBUG] Final line height: {line_height:.1f}pt")
            print(f"🔍 [CERTIFICATE DEBUG] Final total height: {total_height:.1f}pt")
            print(f"🔍 [CERTIFICATE DEBUG] Scope rectangle: {rect}")
            print(f"🔍 [CERTIFICATE DEBUG] Available height: {rect.height:.1f}pt")
            
            if template_type in ["large", "large_eco", "large_nonaccredited"]:
                # Large template: start from top with no margin
                start_y = rect.y0
                print(f"🔍 [CERTIFICATE DEBUG] Large template - Starting from top: y0={rect.y0}")
                
                # Check if text would overflow bottom
                if start_y + total_height > rect.y1:
                    # If overflow, adjust to fit within bounds
                    start_y = rect.y1 - total_height - 2  # 2pt margin from bottom
                    print(f"🔍 [CERTIFICATE DEBUG] ⚠️ Overflow detected! Adjusted start_y to: {start_y:.1f}")
                else:
                    print(f"🔍 [CERTIFICATE DEBUG] ✅ No overflow - using original start_y: {start_y:.1f}")
            else:
                # Standard template: keep current centering logic
                start_y = rect.y0 + (rect.height - total_height) / 2 + line_height/2  # Adjust for baseline
                print(f"🔍 [CERTIFICATE DEBUG] Standard template - Centered positioning: start_y={start_y:.1f}")
            
            print(f"🔍 [CERTIFICATE DEBUG] Final start_y: {start_y:.1f}")
            print(f"🔍 [CERTIFICATE DEBUG] Text will end at: {start_y + total_height:.1f}")
            print(f"🔍 [CERTIFICATE DEBUG] ===== END POSITIONING ANALYSIS =====")
            
            # ✅ SIMPLIFIED: Scope rendering with consistent centering
            current_y = start_y
            
            # Render each line centered for consistent appearance
            for i, line in enumerate(lines):
                if not line.strip():  # Skip empty lines
                    current_y += line_height
                    continue
                
                # Check if this is the last non-empty line
                is_last_line = i == len(lines) - 1 or all(not lines[j].strip() for j in range(i + 1, len(lines)))
                
                if is_last_line:
                    # ✅ LAST LINE: Center align for balanced appearance
                    center_x = (rect.x0 + rect.x1) / 2
                    
                    # ✅ ENHANCED: Use mixed format text rendering for bold detection
                    if '**' in line or '__' in line:
                        # Calculate total width for centering
                        segments = process_bold_text(line)
                        total_width = 0
                        for segment_text, _, _ in segments:
                            if segment_text:
                                font_obj = fitz.Font(fontname="Times-Bold" if "**" in segment_text or "__" in segment_text else "Times-Roman")
                                total_width += font_obj.text_length(segment_text, font_size)
                        
                        start_x = center_x - total_width / 2
                        render_mixed_format_text(page, (start_x, current_y), line, font_size, color)
                    else:
                        # Standard rendering for non-bold text
                        font_obj = fitz.Font(fontname=fontname)
                        text_width = font_obj.text_length(line, font_size)
                        start_x = center_x - text_width / 2
                        
                        page.insert_text(
                            (start_x, current_y),
                            line,
                            fontsize=font_size,
                            fontname=fontname,
                            color=color
                        )
                    # ✅ Safe string handling for debug output
                    safe_line = str(line) if line is not None else ""
                    print(f"🔍 [CERTIFICATE] Last line centered: '{safe_line[:50]}{'...' if len(safe_line) > 50 else ''}'")
                else:
                    # ✅ INTERMEDIATE LINES: Center align for consistency
                    center_x = (rect.x0 + rect.x1) / 2
                    
                    # ✅ ENHANCED: Use mixed format text rendering for bold detection
                    if '**' in line or '__' in line:
                        # Calculate total width for centering
                        segments = process_bold_text(line)
                        total_width = 0
                        for segment_text, _, _ in segments:
                            if segment_text:
                                font_obj = fitz.Font(fontname="Times-Bold" if "**" in segment_text or "__" in segment_text else "Times-Roman")
                                total_width += font_obj.text_length(segment_text, font_size)
                        
                        start_x = center_x - total_width / 2
                        render_mixed_format_text(page, (start_x, current_y), line, font_size, color)
                    else:
                        # Standard rendering for non-bold text
                        font_obj = fitz.Font(fontname=fontname)
                        text_width = font_obj.text_length(line, font_size)
                        start_x = center_x - text_width / 2
                        
                        page.insert_text(
                            (start_x, current_y),
                            line,
                            fontsize=font_size,
                            fontname=fontname,
                            color=color
                        )
                    
                    # ✅ Safe string handling for debug output
                    safe_line = str(line) if line is not None else ""
                    print(f"🔍 [CERTIFICATE] Line {i+1} centered: '{safe_line[:50]}{'...' if len(safe_line) > 50 else ''}'")
                
                # Update current_y consistently for all lines
                # Template-specific line spacing: 1.1 for large/logo, 1.2 for standard
                if template_type in ["large", "large_eco", "large_nonaccredited", "logo", "logo_nonaccredited", "logo_other", "logo_other_nonaccredited"]:
                    current_y += font_size * 1.1  # Tight spacing for large/logo templates
                else:  # standard templates
                    current_y += font_size * 1.2  # Loose spacing for standard templates
            
            print(f"🎯 [SCOPE SUMMARY] Final rendering method: CENTERED")
            print(f"🎯 [SCOPE SUMMARY] Scope field '{field}' processed successfully")
            print(f"🔍 [CERTIFICATE DEBUG] ===== SCOPE ANALYSIS COMPLETE =====")
            print(f"🔍 [CERTIFICATE DEBUG] Final coordinates used: start_y={start_y:.1f}, end_y={start_y + total_height:.1f}")
            print(f"🔍 [CERTIFICATE DEBUG] Space utilization: {((start_y + total_height - rect.y0) / rect.height) * 100:.1f}%")
            print(f"🔍 [CERTIFICATE DEBUG] Remaining space: {rect.y1 - (start_y + total_height):.1f}pt")
            print(f"🔍 [CERTIFICATE DEBUG] ===== END SCOPE ANALYSIS =====")

        
        elif field == "ISO Standard":
            # ✅ ADDED: Preprocess ISO Standard using the mapping
            try:
                # Apply ISO standards mapping to expand short names to full versions
                processed_text = expand_iso_standard(text)
                print(f"🔍 [CERTIFICATE] ISO Standard processed: '{text}' -> '{processed_text}'")
                text = processed_text
            except Exception as e:
                # If preprocessing fails, use text as-is
                print(f"⚠️ [CERTIFICATE] ISO Standard preprocessing failed: {e}, using text as-is: '{text}'")
            
            # Perfect centering for ISO Standard - both horizontal and vertical
            center_x = (rect.x0 + rect.x1) / 2
            center_y = (rect.y0 + rect.y1) / 2 + font_size/3  # Adjust for baseline
            
            # ✅ ENHANCED: Use mixed format text rendering for bold detection
            if '**' in text or '__' in text:
                # Calculate total width for centering
                segments = process_bold_text(text)
                total_width = 0
                for segment_text, _, _ in segments:
                    if segment_text:
                        font_obj = fitz.Font(fontname="Times-Bold" if "**" in segment_text or "__" in segment_text else "Times-Roman")
                        total_width += font_obj.text_length(segment_text, font_size)
                
                start_x = center_x - total_width / 2
                render_mixed_format_text(page, (start_x, center_y), text, font_size, color)
            else:
                # Standard rendering for non-bold text
                font_obj = fitz.Font(fontname=fontname)
                text_width = font_obj.text_length(text, font_size)
                start_x = center_x - text_width / 2
                
                page.insert_text(
                    (start_x, center_y),
                    text,
                    fontsize=font_size,
                    fontname=fontname,
                    color=color
                )
            
            # Print font size for ISO Standard
            print(f"📏 [CERTIFICATE] ISO Standard: {font_size}pt (centered)")
            
            # ✅ MODIFIED: Render certification code below ISO Standard with different coordinates for non-accredited
            try:
                # Check accreditation status - use different coordinates for non-accredited
                accreditation = (values.get("Accreditation") or values.get("accreditation") or "").strip().lower()
                
                # Always draw the certification code, but with different coordinates based on accreditation
                # ✅ ADDED: Defensive check for values dictionary
                if values is None:
                    print(f"⚠️ [CERTIFICATE] Values dictionary is None - skipping certification code")
                else:
                    # Get the certification code for this ISO standard
                    print(f"🔍 [CERTIFICATE] DEBUG: About to get certification code for ISO Standard: '{text}'")
                    
                    # ✅ ADDED: Defensive check for coords dictionary
                    if coords is None:
                        print(f"⚠️ [CERTIFICATE] Coords dictionary is None - skipping certification code")
                    else:
                        print(f"🔍 [CERTIFICATE] DEBUG: Available coordinates keys: {list(coords.keys())}")
                        
                        cert_code = get_iso_standard_code(text)
                        print(f"🔍 [CERTIFICATE] DEBUG: get_iso_standard_code returned: '{cert_code}'")
                        
                        if cert_code:
                            # ✅ NEW: Use different coordinates based on accreditation status
                            if accreditation == "no":
                                # Non-accredited: Move code to the right
                                code_rect = fitz.Rect(335, 757, 390, 762)  # Updated coordinates
                                print(f"🔍 [CERTIFICATE] Non-accredited certificate - using right position")
                            else:
                                # Accredited: Use original position
                                code_rect = coords["certification_code"]  # Original: (253, 757, 285, 762)
                                print(f"🔍 [CERTIFICATE] Accredited certificate - using standard position")
                            
                            print(f"🔍 [CERTIFICATE] DEBUG: Using certification_code coordinates: {code_rect}")
                            
                            # Check if coordinates are within page boundaries
                            page_rect = page.rect
                            print(f"🔍 [CERTIFICATE] DEBUG: Page dimensions: {page_rect}")
                            print(f"🔍 [CERTIFICATE] DEBUG: Code coordinates: x0={code_rect.x0}, y0={code_rect.y0}, x1={code_rect.x1}, y1={code_rect.y1}")
                            
                            if (code_rect.x0 >= 0 and code_rect.y0 >= 0 and 
                                code_rect.x1 <= page_rect.width and code_rect.y1 <= page_rect.height):
                                print(f"🔍 [CERTIFICATE] DEBUG: Coordinates are within page boundaries")
                            else:
                                print(f"⚠️ [CERTIFICATE] DEBUG: WARNING - Coordinates may be outside page boundaries!")
                                print(f"⚠️ [CERTIFICATE] DEBUG: Page width: {page_rect.width}, height: {page_rect.height}")
                                print(f"⚠️ [CERTIFICATE] DEBUG: Code rect: {code_rect}")
                            
                            # Insert certification code with specified font settings
                            print(f"🔍 [CERTIFICATE] DEBUG: About to insert text '{cert_code}' at coordinates ({code_rect.x0}, {code_rect.y0})")
                            
                            # ✅ FIXED: Use reliable font that's always available in PyMuPDF
                            reliable_font = "helv"  # Helvetica - always available in PyMuPDF
                            
                            page.insert_text(
                                (code_rect.x0, code_rect.y0),
                                cert_code,
                                fontsize=5,  # 5pt as specified
                                fontname=reliable_font,  # Use reliable font
                                color=(0, 0, 0)  # Black color
                            )
                            
                            print(f"✅ [CERTIFICATE] Certification code '{cert_code}' rendered at coordinates {code_rect}")
                            print(f"📏 [CERTIFICATE] Certification code: 5pt Times font")
                        else:
                            print(f"⚠️ [CERTIFICATE] No certification code found for ISO Standard: '{text}'")
            except Exception as code_error:
                print(f"⚠️ [CERTIFICATE] Error rendering certification code: {code_error}")
                print(f"⚠️ [CERTIFICATE] Certificate will be generated without certification code")
                import traceback
                print(f"🔍 [CERTIFICATE] Full error traceback: {traceback.format_exc()}")

    # ✅ ADDED: Process Extra Line field
    extra_line_text = values.get("Extra Line", "").strip()
    if extra_line_text:
        print(f"🔍 [CERTIFICATE] Processing Extra Line: '{extra_line_text}'")
        
        # Calculate Extra Line position (0pt gap below scope)
        # Use the same scope_rect that was used for scope rendering
        if template_type in ["large", "large_eco", "large_nonaccredited"]:
            scope_rect = coords["Scope"]  # Single rectangle for large templates
        else:
            # For standard/logo templates, use the stored original coordinates
            if estimated_lines >= 24:
                scope_rect = original_scope_coords["long"]
            else:
                scope_rect = original_scope_coords["short"]
        
        extra_line_y = scope_rect.y1  # 0pt gap - directly below scope
        
        # Create Extra Line rectangle
        extra_line_rect = fitz.Rect(
            scope_rect.x0,      # Same x0 as scope
            extra_line_y,       # 0pt gap - directly below scope
            scope_rect.x1,      # Same x1 as scope  
            extra_line_y + 10   # 10pt height for text
        )
        
        # Render Extra Line text with center alignment and bold font
        if '**' in extra_line_text or '__' in extra_line_text:
            # Use mixed format rendering for bold text with center alignment
            render_mixed_format_text(page, (extra_line_rect.x0, extra_line_rect.y0), extra_line_text, 12, (0, 0, 0), extra_line_rect.width)
        else:
            # Center-aligned bold text rendering
            center_x = (extra_line_rect.x0 + extra_line_rect.x1) / 2
            font_obj = fitz.Font(fontname="Times-Bold")
            text_width = font_obj.text_length(extra_line_text, 12)
            start_x = center_x - text_width / 2
            
            page.insert_text(
                (start_x, extra_line_rect.y0),
                extra_line_text,
                fontsize=12,
                fontname="Times-Bold",
                color=(0, 0, 0)
            )
        
        print(f"🔍 [CERTIFICATE] Extra Line rendered at: {extra_line_rect}")
    else:
        print(f"🔍 [CERTIFICATE] No Extra Line - skipping")

    # ✅ ADDED: Insert logo if available and using logo template
    if logo_image and template_type == "logo":
        try:
            # ✅ ADDED: Defensive check for logo_coords
            if logo_coords is None:
                print(f"⚠️ [CERTIFICATE] Logo_coords is None - skipping logo insertion")
            else:
                # Get logo coordinates from logo_coords
                logo_rect = logo_coords.get("logo")
                if logo_rect:
                    # Use the new shared logo function for better handling
                    insert_logo_into_pdf(page, logo_lookup[logo_filename], logo_rect)
                    print(f"🔍 [CERTIFICATE] Logo inserted successfully using shared logo function")
                else:
                    print("⚠️ [CERTIFICATE] Logo coordinates not found in logo_coords")
        except Exception as logo_insert_error:
            print(f"❌ [CERTIFICATE] Error inserting logo: {logo_insert_error}")

    # ✅ ADDED: Robust return structure - always save and return
    try:
        doc.save(output_pdf_path)
        doc.close()
        
        print(f"[CERTIFICATE] Certificate PDF generated successfully: {output_pdf_path}")
        
        # Return tracking information
        return {
            "success": True,
            "output_path": output_pdf_path,
            "overflow_warnings": overflow_warnings,
            "template_type": template_type
        }
    except Exception as save_error:
        print(f"❌ [CERTIFICATE] Error saving PDF: {save_error}")
        # Still return a result dict even if save fails
        return {
            "success": False,
            "error": f"Failed to save PDF: {save_error}",
            "output_path": output_pdf_path,
            "overflow_warnings": overflow_warnings,
            "template_type": template_type
        }



